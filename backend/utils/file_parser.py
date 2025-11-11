"""
File Parser Utility
Extracts plain text from various file formats (.txt, .md, .pdf, .docx)
"""
import logging
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract plain text from a file based on its extension
    
    Args:
        file_content: Raw file content as bytes
        filename: Original filename (used to determine file type)
        
    Returns:
        Extracted plain text as string
        
    Raises:
        ValueError: If file type is not supported
        Exception: If file parsing fails
    """
    filename_lower = filename.lower()
    
    # Determine file type from extension
    if filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
        return _extract_text_plain(file_content)
    elif filename_lower.endswith('.pdf'):
        return _extract_text_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return _extract_text_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Supported types: .txt, .md, .pdf, .docx")


def _extract_text_plain(file_content: bytes) -> str:
    """Extract text from plain text files (.txt, .md)"""
    try:
        # Try UTF-8 first
        text = file_content.decode('utf-8')
        return text
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1
            text = file_content.decode('latin-1')
            return text
        except UnicodeDecodeError:
            # Last resort: ignore errors
            text = file_content.decode('utf-8', errors='ignore')
            logger.warning("File contained non-UTF-8 characters, some may be lost")
            return text


def _extract_text_pdf(file_content: bytes) -> str:
    """Extract text from PDF files"""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
    
    try:
        text_parts = []
        with pdfplumber.open(BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def _extract_text_docx(file_content: bytes) -> str:
    """Extract text from Word (.docx) files"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for Word document parsing. Install with: pip install python-docx")
    
    try:
        doc = Document(BytesIO(file_content))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

